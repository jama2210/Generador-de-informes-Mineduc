from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

from utils import limpiar_valor, limpiar_nombre_archivo


# -------------------------------------------------------
# ✅ Aplicar tipografía Aptos a todo el documento
# -------------------------------------------------------
def aplicar_tipografia(doc):
    try:
        estilo_normal = doc.styles["Normal"]
        estilo_normal.font.name = "Aptos"
        estilo_normal.font.size = Pt(11)

        # Aplicar Aptos a todos los estilos posibles
        for estilo in doc.styles:
            try:
                estilo.font.name = "Aptos"
            except:
                pass

    except Exception as e:
        print("Error aplicando tipografía:", e)


# -------------------------------------------------------
# ✅ Función principal para generar informes
# -------------------------------------------------------
def generar_informes(df, carpeta, barra, estado, modo):

    grupos = df.groupby(["Indique su región", "Deprov", "Tipo Asesoría"])
    total = len(grupos)
    contador = 0

    for (region, deprov, modalidad), datos_grupo in grupos:

        # Crear subcarpetas por Región / Deprov / Modalidad
        base_sub = os.path.join(
            carpeta,
            limpiar_nombre_archivo(region),
            limpiar_nombre_archivo(deprov),
            limpiar_nombre_archivo(modalidad)
        )
        os.makedirs(base_sub, exist_ok=True)

        # ----------------------------------------------
        # ✅ VARIANTE A — Informe grupal por R-D-M
        # ----------------------------------------------
        if "Variante A" in modo:

            doc = Document()
            aplicar_tipografia(doc)

            doc.add_picture("logo_mineduc.png", width=Cm(4))
            titulo = doc.add_heading(
                "Informe de Planificación de Asesoría Ministerial",
                level=0
            )
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            tabla = doc.add_table(rows=3, cols=2)
            tabla.style = "Table Grid"
            tabla.cell(0,0).text = "Región"
            tabla.cell(0,1).text = limpiar_valor(region)

            tabla.cell(1,0).text = "DEPROV"
            tabla.cell(1,1).text = limpiar_valor(deprov)

            tabla.cell(2,0).text = "Modalidad"
            tabla.cell(2,1).text = limpiar_valor(modalidad)

            doc.add_heading(
                f"Profesionales incluidos ({len(datos_grupo)})",
                level=1
            )

            datos_grupo = datos_grupo.sort_values("Nombre")

            for _, row in datos_grupo.iterrows():

                doc.add_heading(
                    limpiar_valor(row.get("Nombre")),
                    level=2
                )

                tabla_p = doc.add_table(rows=0, cols=2)
                tabla_p.style = "Table Grid"

                columnas_ignorar = [
                    "ID",
                    "Hora de inicio",
                    "Hora de finalización"
                ]

                for col in datos_grupo.columns:
                    if col in columnas_ignorar:
                        continue

                    fila = tabla_p.add_row()
                    fila.cells[0].text = col
                    fila.cells[1].text = limpiar_valor(row[col])

                doc.add_paragraph()

            nombre_archivo = (
                f"Informe_{limpiar_nombre_archivo(region)}_"
                f"{limpiar_nombre_archivo(deprov)}_"
                f"{limpiar_nombre_archivo(modalidad)}.docx"
            )

            ruta = os.path.join(base_sub, nombre_archivo)
            doc.save(ruta)

        # ----------------------------------------------
        # ✅ VARIANTE B — Un informe por profesional
        # ----------------------------------------------
        else:

            # Agrupar cada profesional
            for nombre_prof, datos_persona in datos_grupo.groupby("Nombre"):

                doc = Document()
                aplicar_tipografia(doc)

                doc.add_picture("logo_mineduc.png", width=Cm(4))
                titulo = doc.add_heading(
                    "Informe Individual de Asesoría MINEDUC",
                    level=0
                )
                titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()

                # Encabezado del informe
                tabla = doc.add_table(rows=3, cols=2)
                tabla.style = "Table Grid"
                tabla.cell(0,0).text = "Región"
                tabla.cell(0,1).text = limpiar_valor(region)

                tabla.cell(1,0).text = "DEPROV"
                tabla.cell(1,1).text = limpiar_valor(deprov)

                tabla.cell(2,0).text = "Modalidad"
                tabla.cell(2,1).text = limpiar_valor(modalidad)

                doc.add_heading(limpiar_valor(nombre_prof), level=1)

                # Ordenar registros por Nombre Asesoría si existe
                if "Nombre Asesoría" in datos_persona.columns:
                    datos_persona = datos_persona.sort_values("Nombre RBD, RED, Sostenedor")

                # Iterar cada uno de los registros del profesional (MUY IMPORTANTE)
                for _, fila_persona in datos_persona.iterrows():

                    nombre_asesoria = limpiar_valor(
                        fila_persona.get("Nombre RBD, RED, Sostenedor", "Registro")
                    )

                    doc.add_heading(
                        f"Registro asociado a: {nombre_asesoria}",
                        level=2
                    )

                    tabla_p = doc.add_table(rows=0, cols=2)
                    tabla_p.style = "Table Grid"

                    columnas_ignorar = [
                        "ID",
                        "Hora de inicio",
                        "Hora de finalización"
                    ]

                    for col in datos_persona.columns:

                        if col in columnas_ignorar:
                            continue

                        fila = tabla_p.add_row()
                        fila.cells[0].text = col
                        fila.cells[1].text = limpiar_valor(fila_persona[col])

                    doc.add_paragraph()

                # Guardar archivo individual
                nombre_archivo = (
                    f"Informe_{limpiar_nombre_archivo(region)}_"
                    f"{limpiar_nombre_archivo(deprov)}_"
                    f"{limpiar_nombre_archivo(modalidad)}_"
                    f"{limpiar_nombre_archivo(nombre_prof)}.docx"
                )

                ruta = os.path.join(base_sub, nombre_archivo)
                doc.save(ruta)

        # ----------------------------------------------
        # ✅ Barra de progreso
        # ----------------------------------------------
        contador += 1
        barra.progress(contador / total)
        estado.text(f"Generando informe {contador} de {total}")